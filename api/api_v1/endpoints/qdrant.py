import os
import uuid
import qdrant_client

from sentence_transformers import SentenceTransformer
from typing import Union
from tqdm import tqdm
from fastapi import APIRouter, Body, Depends, HTTPException, UploadFile, Request
from docx import Document
from pydantic import BaseModel
from qdrant_client.http.models import Distance, VectorParams


from langchain.embeddings.sentence_transformer import SentenceTransformerEmbeddings, HuggingFaceEmbeddings

from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores.qdrant import Qdrant
from langchain.document_loaders import TextLoader
from langchain.schema import document


CHUNK_SIZE = 2**20  # 1MB
FILES_PATH = "./temp"

router = APIRouter()
embeddings = SentenceTransformerEmbeddings(model_name="snunlp/KR-SBERT-V40K-klueNLI-augSTS")
# embeddings = HuggingFaceEmbeddings( model_name="sentence-transformers/all-mpnet-base-v2")
client = qdrant_client.QdrantClient(
    host="localhost",
    # url="http://localhost:6333/",
    # port=6333
    # api_key="<qdrant-api-key>", # For Qdrant Cloud, None for local instance
)

# client.recreate_collection

# client.recreate_collection('qdrant_test', vectors_config=VectorParams(size=1000, distance=Distance.DOT))

db = Qdrant(
    client=client, 
    collection_name="qdrant_test",
    embeddings=embeddings,
)

class Source(BaseModel):
    name: str
    price: float
    is_offer: Union[bool, None] = None


async def chunked_copy(src, dst):
    await src.seek(0)        
    with open(dst, "wb") as buffer:
        while True:
            contents = await src.read(CHUNK_SIZE)
            # print(contents)
            if not contents:
                break
            buffer.write(contents)


@router.get("/qdrant/{query}")
def read_source(
    query: str,
    # count: int,
    distance: float
):
    # db = Qdrant(
    #     embedding_function,
    #     connection_args={"host": "127.0.0.1", "port": "19530"},
    #     collection_name="collection_1",
    # )
    items = db.similarity_search_with_score(query=query, filter={"user": "keonlee.gun@gmail.com"}, )
    res = []
    
    for item in items:
        if(item[1] < distance):
            res.append(item)    
    
    return res

@router.post("/qdrant")
async def create_source(
    file: UploadFile,
    req: Request    
):
    doc = Document(file.file)
    # ids = []
    # metas = []
    # text = []
    texts = []
    
    for idx, para in enumerate(doc.paragraphs):
        sentence = para.text.strip()
        data = document.Document(page_content=sentence, metadata={"user": "keonlee.gun@gmail.com"})
        texts.append(data)
    
    db.add_documents(texts)
        
        
    #     sentence = para.text.strip()
    #     text.append(sentence)
    #     ids.append(str(idx))
    #     metas.append({"user": "keonlee.gun@gmail.com"})

    # db.add_texts(texts=text, ids=ids, metadatas=metas)

    return
