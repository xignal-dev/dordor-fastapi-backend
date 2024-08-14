import os
import uuid
import sys
import chromadb
import datetime, time, json
import string, random
from cryptography.fernet import Fernet
import hashlib
from fastapi.security import OAuth2PasswordBearer, APIKeyHeader

# from sentence_transformers import SentenceTransformer
from typing import (IO, Dict, Iterable, Iterator, Mapping, Optional, Tuple, Union, Annotated)
from tqdm import tqdm
from fastapi import APIRouter, Body, Depends, HTTPException, UploadFile, Request
from docx import Document
from pydantic import BaseModel

from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain.vectorstores import chroma
from langchain.document_loaders import Docx2txtLoader
from langchain.schema import document

from module import file_cryptography

CHUNK_SIZE = 2**20  # 1MB
# FILES_PATH = os.environ.get("FILES_PATH")

# load .env
# load_dotenv()

# 환경변수
FILES_PATH = os.environ.get("FILES_PATH")
D_ORIGINALS = os.environ.get("D_ORIGINALS")
D_FORMAT = os.environ.get("D_FORMAT")
D_EDITED = os.environ.get("D_EDITED")
D_DOC = os.environ.get("D_DOC")
D_EXCEL = os.environ.get("D_EXCEL")
D_GLOSSARY = os.environ.get("D_GLOSSARY")

ORIGINALS_PATH = os.path.join(FILES_PATH, D_ORIGINALS)
FORMAT_PATH = os.path.join(FILES_PATH, D_FORMAT)
EDITED_PATH = os.path.join(FILES_PATH, D_EDITED)
DOC_PATH = os.path.join(FILES_PATH, D_DOC)
EXCEL_PATH = os.path.join(FILES_PATH, D_EXCEL)
GLOSSARY_PATH = os.path.join(FILES_PATH, D_GLOSSARY)

# ORIGINALS_PATH = os.path.join('FILES_PATH', 'D_ORIGINALS')
# FORMAT_PATH = os.path.join('FILES_PATH', 'D_FORMAT')
# EDITED_PATH = os.path.join('FILES_PATH', 'D_EDITED')
# DOC_PATH = os.path.join('FILES_PATH', 'D_DOC')
# EXCEL_PATH = os.path.join('FILES_PATH', 'D_EXCEL')
# GLOSSARY_PATH = os.path.join('FILES_PATH', 'D_GLOSSARY')
# oauth2 = OAuth2PasswordBearer(tokenUrl="token")
apikey = APIKeyHeader(name="token")

router = APIRouter()
# model = SentenceTransformer('all-MiniLM-L6-v2')
# model = SentenceTransformer('snunlp/KR-SBERT-V40K-klueNLI-augSTS')
kr_embedding_function = SentenceTransformerEmbeddings(model_name="ddobokki/klue-roberta-small-nli-sts")
en_embedding_function = SentenceTransformerEmbeddings(model_name="sentence-transformers/all-MiniLM-L12-v2")
# en_embedding_function = SentenceTransformerEmbeddings(model_name="microsoft/Multilingual-MiniLM-L12-H384")
# en_embedding_function = SentenceTransformerEmbeddings(model_name="microsoft/MiniLM-L12-H384-uncased")

chromaClient = chromadb.HttpClient(host="172.17.0.1", port=18001)
# chromaClient = chromadb.HttpClient(host="localhost", port=18001)

# colList = chromaClient.list_collections()
# print(chromaClient.list_collections())

# collection = chromaClient.list_collections().index(0)
# if(colList.__len__() < 1):
# kr_collection = chromaClient.get_or_create_collection(name="kr_src_collection", )
# en_collection = chromaClient.get_or_create_collection(name="en_src_collection", )
kr_collection = chromaClient.get_or_create_collection(name="kr_src_collection", embedding_function=kr_embedding_function)
en_collection = chromaClient.get_or_create_collection(name="en_src_collection", embedding_function=en_embedding_function)
config = chromaClient.get_or_create_collection(name="config_collection", embedding_function=kr_embedding_function)


kr_db = chroma.Chroma(
    client=chromaClient,
    collection_name="kr_src_collection",
    embedding_function=kr_embedding_function, 
)

en_db = chroma.Chroma(
    client=chromaClient,
    collection_name="en_src_collection",
    embedding_function=en_embedding_function, 
)

config_db = chroma.Chroma(
    client=chromaClient,
    collection_name="config_collection",
    embedding_function=kr_embedding_function, 
)

# db = chroma.Chroma(client=chromaClient, collection_name="test_collection", embedding_function=embedding_function)

class TMData(BaseModel):
    text: str | None = None
    user: str | None = None
    docType: str | None = None
    docName: str | None = None
    language: str | None = None
    status: str | None = None
    translated: str | None = None
StrPath = Union[str, 'os.PathLike[str]']

async def chunked_copy(src, dst):
    await src.seek(0)        
    with open(dst, "wb") as buffer:
        while True:
            contents = await src.read(CHUNK_SIZE)
            # print(contents)
            if not contents:
                break
            buffer.write(contents)

def getDB(lang):
    db = kr_db
        
    if lang == 'en_ko':
        db = en_db
    elif lang == 'ko_en':
        db = kr_db
        
    return db

def initDB(lang):
        
    if(lang == 'ko_en'):
        global kr_db
        chromaClient.get_or_create_collection(name="kr_src_collection")
        kr_db = chroma.Chroma(
            client=chromaClient,
            collection_name="kr_src_collection",
            embedding_function=kr_embedding_function, 
        )
    elif(lang == 'en_ko'):
        global en_db
        chromaClient.get_or_create_collection(name="en_src_collection")
        en_db = chroma.Chroma(
            client=chromaClient,
            collection_name="en_src_collection",
            embedding_function=kr_embedding_function, 
        )
        
# 원문 데이터 읽어오는 타이밍
# 번역 데이터 읽어오는 타이밍: 처음에는 원문과 함께, 번역 작업 완료 시에는 번역 데이터만 (혹은 원문과 함께?) 읽어온 후 최종 업데이트함
# langchain - chroma db의 각 레코드 형식은 아래와 같이 각각의 타입으로 나뉘어져 있다.
# ids: 각 레코드의 고유한 ID 값
# texts: 유사도 검색 대상의 텍스트 데이터 값
# metadatas: 기타 데이터 값
# TM의 ID는 ids에, 원문은 texts에, 번역 데이터, 사용자 정보 및 기타 필요한 정보는 metadatas에 입력한다. 
# 중복된 TM 원문 입력 시, 즉 새로 입력되는 TM이 기존 TM과의 유사도가 같은 경우, 원문이 기존 TM과 중복되는 새로운 TM을 추가한다.
# 기존의 TM 값을 metadata 조건에 맞게 모두 요청하는 기능 
# 기존의 TM 값을 삭제하는 기능
# 기존의 TM 값을 수정하는 기능
# 유사도 비교 기준 값
# metadatas에 추가할 정보: user_email, created_at, modified_at, doc_name, doc_type, language, status, translated_text
import asyncio, time

async def say_after(delay, what):
    await asyncio.sleep(delay)
    print(what)

def encode_test(key, clientIp: str):
    
    message = clientIp

    fernet = Fernet(key)
    
    encMessage = fernet.encrypt(message.encode())
    
    print("original string: ", clientIp)
    print("encrypted string: ", encMessage)
    
    decMessage = fernet.decrypt(encMessage).decode()
    
    print("decrypted string: ", decMessage)

    return encMessage

def checkAdminAuth(token):
    res = True
    auth = config_db.get(where={"email": 'admin'})
    print(auth)
    if(auth['metadatas'] == [] or auth['metadatas'][0]['apikey'] != token):
        res = False

    return res

def checkUserAuth(token):
    res = True
    auth = config_db.get(where={"apikey": token})
    print(auth)
    if(auth['metadatas'] == [] or auth['metadatas'][0]['isValid'] != True):
        res = False

    return res

class UserInfo(BaseModel):
    user: str | None = None
    password: str | None = None

@router.post("/user", tags=["users"], name="TM 서비스 사용자 계정 생성")
def create_user(userInfo: UserInfo):
    ids = []
    metas = []
    text = []
    
    role = "user"
    if(userInfo.user == "admin"):
        role = "admin"
    
    item = config_db.get(where={"email": userInfo.user})
    if(len(item['metadatas']) > 0):
        print(item)
        
        if not 'role' in item['metadatas'][0]:
            item['metadatas'][0]['role'] = role
            data = document.Document(page_content= item['documents'][0], metadata=item['metadatas'][0])
            config_db.update_document(item['ids'][0], data)
            print(item['metadatas'][0])
        
        item = {"message": userInfo.user + ' is already exist.'}
    else:
        item = {}
        text.append(userInfo.user)
        metas.append({"email": userInfo.user, "apikey": "", "password": userInfo.password, "isValid": False, "role": role})
        config_db.add_texts(texts=text, metadatas=metas)
        
        print(metas)
        item = {"message": userInfo.user + ' is created.'}
    return item

class UpdateUserInfo(BaseModel):
    user: str | None = None
    oldPassword: str | None = None
    newPassword: str | None = None
    
@router.patch("/user/password", tags=["users"], name="TM 서비스 사용자 패스워드 변경")
def update_password(userInfo: UpdateUserInfo):
    res = {"message": userInfo.user + " is not exist."}
        
    item = config_db.get(where={"email": userInfo.user})
    
    # print(item['metadatas'][0])   
    if(len(item['metadatas']) > 0):

        if not 'password' in item['metadatas'][0]:
            item['metadatas'][0]['password'] = userInfo.oldPassword
                
        if(userInfo.oldPassword == item['metadatas'][0]['password']):
        
            id = item['ids'][0]
            doc = item['documents'][0]
            metadata = {"email": item['metadatas'][0]['email'], "apikey": item['metadatas'][0]['apikey'], "password": userInfo.newPassword, "isValid": item['metadatas'][0]['isValid']}
            data = document.Document(page_content=doc, metadata=metadata)
            config_db.update_document(id, data)
            print(metadata)
            res = {"message": "Update new password"}
        else:
            res = {"message": "Wrong old password"}
        
    return res

@router.post("/user/apikey", tags=["users"], name="TM 서비스 하용자 API Key 발급")
def get_or_generate_api_key(
    userInfo: UserInfo,
    newgen: int | None = None):
    
    item = config_db.get(where={"email": userInfo.user})
        
    if(len(item['metadatas']) > 0):
        print(item)
        
        if(userInfo.password != item['metadatas'][0]['password']):
            return {"message": "Wrong password"}
        
        if(item['metadatas'][0]['apikey'] == '' or newgen == 1):
            now = datetime.datetime.now()
            key = hashlib.sha256(userInfo.user.encode() + str(now).encode()).hexdigest()
    
            id = item['ids'][0]
            doc = item['documents'][0]
            metadata = {"email": item['metadatas'][0]['email'], "apikey": key, "password": item['metadatas'][0]['password'], "isValid": item['metadatas'][0]['isValid']}
            data = document.Document(page_content=doc, metadata=metadata)
            config_db.update_document(id, data)
            item = {"email": metadata['email'], "apikey": key, "isValid": metadata['isValid']}
        else:
            item = {"email": item['metadatas'][0]['email'], "apikey": item['metadatas'][0]["apikey"], "isValid": item['metadatas'][0]['isValid']}
    else:
        return {"message": userInfo.user + " is not exist"}
    
    return item

# @router.patch("/user/apikey")
# def update_api_key(userInfo: UserInfo):
    
#     text = []
#     metas = []
    
#     item = config_db.get(where={"email": userInfo.user})
    
#     now = datetime.datetime.now()
    
#     if(len(item['metadatas']) > 0):
        
#         if(userInfo.password != item['metadatas'][0]['password']):
#             return {"message": "Wrong password"}
    
#         key = hashlib.sha256(userInfo.user.encode() + str(now).encode()).hexdigest()
#         id = item['ids'][0]
#         doc = item['documents'][0]
#         metadata = {"email": item['metadatas'][0]['email'], "apikey": key, "password": item['metadatas'][0]['password'], "isValid": item['metadatas'][0]['isValid']}
#         # metadata = {"email": userInfo.user, "apikey": key, "isValid": True}
#         data = document.Document(page_content=doc, metadata=metadata)
#         config_db.update_document(id, data)
#         print(metadata)
#     else:
#         return {"message": userInfo.user + " is not exist"}
    
#     return metadata

# ee263b0f380713589f9f9f3fbcbda2b06ca0950cd69192e8d44b1ae4f8d90583
@router.get("/admin/user", tags=["admin"], name="TM 서비스 사용자 계정 리스트 조회")
def get_user_list(
    token: Annotated[str, Depends(apikey)]):
    
    if(checkAdminAuth(token) != True):
        return {}
    
    items = config_db.get()
    # if(len(items['metadatas']) > 0):
    res = []
    # idx = 0
    for item in items['metadatas']:
        # print(item)
        if(item['email'] == 'admin'):
            continue
        
        res.append({'email': item['email'], 'isValid': item['isValid']})

    return res

@router.post("/admin", tags=["admin"], name="TM 서비스 관리자(admin) 계정 생성")
def create_admin(
    password: str, 
    ):
    
    res = {}
    
    item = config_db.get(where={"email": 'admin'})
    
    if(len(item['metadatas']) > 0):
        res = {"message": "admin already exists."}
    else:
        metas = []
        text = []
        text.append("admin")
        metas.append({"email": "admin", "apikey": "", "password": password, "isValid": True, "role": "admin"})
        config_db.add_texts(texts=text, metadatas=metas)
        
        print(metas)
        res = { "email": "admin", "isValid": "true" }
    
    return res


@router.patch("/admin/user", tags=["admin"], name="ㅅTM 서비스 사용자 계정 활성/비활성화")
def change_validation(
    token: Annotated[str, Depends(apikey)],
    user: str, 
    isValid: bool | None = None):
    
    if(checkAdminAuth(token) != True):
        return {}
    
    res = {"message": user + " does not exist."}
    item = config_db.get(where={"email": user})
    
    # print(item['metadatas'][0])   
    if(len(item['metadatas']) > 0):
        id = item['ids'][0]
        doc = item['documents'][0]
        metadata = {"email": item['metadatas'][0]['email'], "apikey": item['metadatas'][0]['apikey'], "password": item['metadatas'][0]['password'], "isValid": isValid}
        data = document.Document(page_content=doc, metadata=metadata)
        config_db.update_document(id, data)
        # print(metadata)
        res = {"message": user + " validation is set as " + str(isValid)}
        
    return res

@router.delete("/admin/user", tags=["admin"], name="TM 사용자 계정 제거")
def delete_user(
    token: Annotated[str, Depends(apikey)],
    user: str):
    
    if(checkAdminAuth(token) != True):
        return {}
    
    res = {"message": user + " is not exist."}
        
    where = { "email": user}
    item = config_db.get(where=where)
    
    if(len(item['metadatas']) > 0):
        config_db.delete(ids=item['ids'])
        res = {"message": user + " is deleted"}

    return res

# @router.get("/apikey")
# def get_api_key(
#     # token: Annotated[str, Depends(apikey)], 
#     email: str):
    
#     # if(checkAuth(token) != True):
#     #     return {}
    
#     item = config_db.get(where={"email": email})
#     if(len(item['metadatas']) > 0):
#         print(item)
#         item = item['metadatas'][0]
#     else:
#         item = {}
#     #     text.append(email)
#     #     key = hashlib.sha256(email.encode()).hexdigest()

#     #     # ids.append(timestamp + '_' + str(idx))
#     #     metas.append({"email": email, "apikey": key, "isValid": "true"})
#     #     config_db.add_texts(texts=text, metadatas=metas)
        
#     #     print(metas)
#     #     item = metas
#     return item

@router.get("/chroma/test", tags=["chroma"], name="크로마 테스트 API")
async def get_test(req: Request):
    
    # task = asyncio.create_task(say_after(1, 'hello'))
    # task
    encMessage = {"message": "Wrong old password"}
    print(encMessage)

    return encMessage


@router.get("/chroma/{lang}/{key}", tags=["chroma"], name="주어진 언어와 키/값 조건에 따른 TM 데이터 가져오기")
def get_by_key(
    token: Annotated[str, Depends(apikey)],
    lang: str,
    key: str,
    value: str):
    
    if(checkUserAuth(token) != True):
        return {}
    
    db = getDB(lang=lang)    
    where = { key: value }
    items = db.get(where=where)
    
    return items
    
class QueryInfo(BaseModel):
    query: str | None = None
    
@router.post("/chroma/{lang}/similarity", tags=["chroma"], name="주어진 언어와 문장에 대한 TM 데이터 유사도 검색")
def similarity_search(
    token: Annotated[str, Depends(apikey)],
    lang: str,
    queryInfo: QueryInfo,
    user: str,
    distance: float
):
    
    if(checkUserAuth(token) != True):
        return {}

    
    db = getDB(lang=lang)
    # items = db.similarity_search_with_score(query=query, filter={"user": "keonlee.gun@gmail.com"}, )
    items = db.similarity_search_with_score(query=queryInfo.query, filter={"user": user}, k=1 )

    # res = db.__query_collection(where={"user": "keonlee.gun@gmail.com"}, query_texts=[query])
    # res = db.similarity_search(query, filter={"user": "keonlee.gun@gmail.com"}, )
    res = []
    
    for item in items:
        print(item)
        if(item[1] < distance):
            res.append(item)    
    
    return res

# @router.get("/chroma/{lang}/similarity/")
# def similarity_search_from_file(
#     lang: str,
#     user: str,
#     # count: int,
#     distance: float
# ):
#     db = getDB(lang=lang)
#     fpath = os.path.expanduser('/Users/keonlee/dev_home/knb/Tlexibeta-TM/temp/originals/contract.docx')
#     fb = open(fpath, 'rb')
#     doc = Document(fb)
    
#     allItems = []
    
#     for para in doc.paragraphs:
#         query = para.text.strip()
#         items = db.similarity_search_with_score(query=query, filter={"user": user}, )
#         allItems.append(items)
        
#     res = []
    
#     for items in allItems:
#         temp = []
#         for item in items:
#             if(item[1] < distance):
#                 temp.append(item)
        
#         res.append(temp)
    
#     return res

@router.delete("/chroma/{lang}", tags=["chroma"], name="주어진 언어에 맞는 모든 TM 데이터 삭제")
async def delete_all_data(
    token: Annotated[str, Depends(apikey)],
    lang: str):
    
    if(checkUserAuth(token) != True):
        return {}

    db = getDB(lang=lang)    
    db.delete_collection()
    
    initDB(lang)

    return "OK"

@router.delete("/chroma/{lang}/{key}", tags=["chroma"], name="주어진 언어와 키/값 조건에 맞는 TM 데이터 삭제")
async def delete_by_id(
    token: Annotated[str, Depends(apikey)],
    lang: str,
    key: str,
    value: str):
    
    if(checkUserAuth(token) != True):
        return {}
    
    db = getDB(lang=lang)
    
    where = { key: value }    
    items = db.get(where=where)
    db.delete(ids=items['ids'])
    return "OK"

@router.patch("/chroma/{lang}/{id}", tags=["chroma"], name="주어진 언어와 ID에 대한 TM 데이터를 수정")
async def patch_by_id(
    token: Annotated[str, Depends(apikey)],
    lang: str,
    id: str,
    user: str | None = None,
    docType: str | None = None,
    docName: str | None = None,
    language: str | None = None,
    status: str | None = None,
):
    
    if(checkUserAuth(token) != True):
        return {}

    res = "NO_ID"
    db = getDB(lang=lang)
    # metas.append({"user": "keonlee.gun@gmail.com", "created_at": now, "modified_at": now, "doc_name": file.filename, "doc_type": "landr", "language": "En", "status": "enable", "id": timestamp + '_' + str(idx)})
    now = str(datetime.datetime.now())
    tmData = db.get(id)
    
    if(tmData):
        meta = tmData['metadatas'][0]
        doc = tmData['documents'][0]
        # doc = text if text else db.get(id)['documents'][0]    
        # print(meta)
        
        metadata={"user": user if user else meta['user'], "created_at": meta['created_at'], "modified_at": now, "doc_name": docName if docName else meta['doc_name'], "doc_type": docType if docType else meta['doc_type'], "language": language if language else meta['language'], "status": status if status else meta['status'], "id": meta['id']}
        # "translated_text": translated
        
        data = document.Document(page_content=doc, metadata=metadata)    
        db.update_document(id, data)
        res = "OK"        
    
    return res

@router.post("/chroma/{lang}", tags=["chroma"], name="주어진 언어에 맞는 TM 데이터를 입력")
async def create_tm_data(
    token: Annotated[str, Depends(apikey)],
    lang: str,
    tmData: TMData):
    
    if(checkUserAuth(token) != True):
        return {}

    
    db = getDB(lang=lang)
    ids = []
    metas = []
    text = []
    idx = 0
    timestamp = str(time.time())
    now = str(datetime.datetime.now())
    
    text.append(tmData.text)
    ids.append(timestamp + '_' + str(idx))
    metas.append({"user": tmData.user, "created_at": now, "modified_at": now, "doc_name": tmData.docName, "doc_type": tmData.docType, "language": lang, "status": tmData.status, "translated": tmData.translated, "id": timestamp + '_' + str(idx)})
    # "translated_text": translated
    
    db.add_texts(texts=text, ids=ids, metadatas=metas)
    return "OK"

@router.post("/chroma/{lang}/file/{id}", tags=["chroma"], name="주어진 언어에 맞는 TM 데이터를 번역 파일로부터 입력")
async def import_tm_file_by_idx(
    token: Annotated[str, Depends(apikey)],
    lang: str,
    user: str,
    id: str,
    docType: str,
    docName: str,
    req: Request    
):
    if(checkUserAuth(token) != True):
        return {}

    db = getDB(lang=lang)

    fullpath2 = os.path.join(EDITED_PATH + "/", id)

    file_cryptography.file_decryption(encryption_file_name=fullpath2, decryption_file_name=fullpath2)
    fb2 = open(f"/{EDITED_PATH}/{id}", 'rb')

    doc2 = json.load(fb2)
    doc2 = json.loads(doc2)

    file_cryptography.file_encryption(original_file_name=fullpath2, encryption_file_name=fullpath2)
    
    ids = []
    metas = []
    text = []
    nIdx = 0
    timestamp = str(time.time())
    now = str(datetime.datetime.now())
    
    for para in doc2["paragraphs"]:
        sentence = para['src']
        translated = para['tgt']
        text.append(sentence)
        ids.append(timestamp + '_' + str(nIdx))
        metas.append({"user": user, "created_at": now, "modified_at": now, "doc_name": docName, "doc_type": docType, "language": lang, "status": "enable", "id": timestamp + '_' + str(nIdx), "translated_text": translated})
        nIdx = nIdx + 1
    
    res = "NO_SRC"
    
    if(nIdx > 1):
        db.add_texts(texts=text, ids=ids, metadatas=metas)
        res = "OK"
        
    return res

# @router.post("/chroma/{lang}/file")
# async def import_tm_file(
#     token: Annotated[str, Depends(apikey)],
#     lang: str,
#     req: Request):
    
#     if(checkUserAuth(token) != True):
#         return {}

    
#     db = getDB(lang=lang)
#     fpath = os.path.expanduser('/Users/keonlee/dev_home/knb/Tlexibeta-TM/temp/originals/contract.docx')
#     fb = open(fpath, 'rb')
#     doc = Document(fb)
    
#     ids = []
#     metas = []
#     text = []
#     idx = 0
#     timestamp = str(time.time())
#     now = str(datetime.datetime.now())
    
#     for para in doc.paragraphs:
#         sentence = para.text.strip()
#         if(len(sentence) == 0):
#             continue        

#         text.append(sentence)
#         ids.append(timestamp + '_' + str(idx))
#         metas.append({"user": "keonlee.gun@gmail.com", "created_at": now, "modified_at": now, "doc_name": 'contract.docx', "doc_type": "landr", "language": lang, "status": "enable", "id": timestamp + '_' + str(idx)})
#         idx = idx + 1

#     res = "NO_SRC"
    
#     if(idx > 1):
#         db.add_texts(texts=text, ids=ids, metadatas=metas)
#         res = "OK"
        
#     return res

