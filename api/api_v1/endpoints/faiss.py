from typing import Union
from tqdm import tqdm
from fastapi import APIRouter, Body, Depends, HTTPException, UploadFile, Request
from docx import Document
from pydantic import BaseModel

from langchain.embeddings.sentence_transformer import SentenceTransformerEmbeddings

from langchain.vectorstores.faiss import FAISS
from langchain.docstore.in_memory import InMemoryDocstore
# from langchain.docstore import

CHUNK_SIZE = 2**20  # 1MB
FILES_PATH = "./temp"

router = APIRouter()
embedding_function = SentenceTransformerEmbeddings(model_name="snunlp/KR-SBERT-V40K-klueNLI-augSTS")

db = FAISS(embedding_function=embedding_function, index='faiss_local', docstore=InMemoryDocstore({}), index_to_docstore_id={})


class Source(BaseModel):
    name: str
    price: float
    is_offer: Union[bool, None] = None


async def chunked_copy(src, dst):
    await src.seek(0)
    with open(dst, "wb") as buffer:
        while True:
            contents = await src.read(CHUNK_SIZE)
            # print(contents)
            if not contents:
                break
            buffer.write(contents)


@router.get("/faiss/{query}")
def read_source(
    query: str,
    # count: int,
    distance: float
):

    items = db.similarity_search_with_score(query=query, filter={"user": "keonlee.gun@gmail.com"}, )
    res = []
    
    for item in items:
        if(item[1] < distance):
            res.append(item)    
    
    return res

@router.post("/faiss")
async def create_source(
    file: UploadFile,
    req: Request    
):
    doc = Document(file.file)
    ids = []
    metas = []
    text = []
    # texts = []

    
    for idx, para in enumerate(doc.paragraphs):
        sentence = para.text.strip()
        text.append(sentence)
        ids.append(str(idx))
        metas.append({"user": "keonlee.gun@gmail.com"})
        # sentence = para.text.strip()
        # data = document.Document(page_content=sentence, metadata={"user": "keonlee.gun@gmail.com"})
        # texts.append(data)
    
    # db.add_documents(texts)
    db.add_texts(texts=text, ids=ids, metadatas=metas)

    return
